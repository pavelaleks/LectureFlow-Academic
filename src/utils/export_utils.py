"""
Export utilities for lectures and related content.
"""
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_markdown(content: str, file_path: str | Path) -> None:
    """Export content as markdown file."""
    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)


def export_docx(content: str, file_path: str | Path, title: str = "Lecture") -> None:
    """Export content as DOCX file (simple implementation)."""
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        # Check if it's a heading (starts with #)
        if para_text.startswith('#'):
            level = len(para_text) - len(para_text.lstrip('#'))
            heading_text = para_text.lstrip('#').strip()
            doc.add_heading(heading_text, level=min(level, 3))
        else:
            # Regular paragraph
            p = doc.add_paragraph(para_text)
            p.style.font.size = Pt(11)
    
    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))

