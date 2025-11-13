"""
DOCX export utilities for lectures.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import re


def md_to_docx_paragraph(document: Document, text: str):
    """
    Converts simple Markdown (**bold**, *italic*) into docx formatted paragraphs.
    Supports Unicode and handles escaped characters properly.
    
    Args:
        document: Word document instance
        text: Text with Markdown formatting
    """
    paragraph = document.add_paragraph()
    
    # Process markdown formatting: **bold** and *italic*
    # Pattern: matches **bold** or *italic*
    # Use non-greedy matching and ensure we capture complete markdown blocks
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Check for bold (**text**)
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        # Check for italic (*text*) - but not **bold** or ***both***
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            # Skip if it's part of bold or triple asterisk
            if not (part.startswith("***") or part.endswith("***")):
                italic_text = part[1:-1]
                run = paragraph.add_run(italic_text)
                run.italic = True
            else:
                # Regular text - preserve as is
                paragraph.add_run(part)
        else:
            # Regular text - preserve as is (Unicode safe)
            # No escaping needed, python-docx handles Unicode correctly
            paragraph.add_run(part)


def export_lecture_to_docx(
    title: str,
    subtitle: str,
    keywords: list,
    lecture_text: str,
    bibliography: str | None = None,
    file_path: str | Path = None
) -> None:
    """
    Create a fully formatted Word (.docx) file from the generated lecture.
    
    Args:
        title: Lecture title
        subtitle: Lecture subtitle (optional)
        keywords: List of keywords
        lecture_text: Main lecture text (can contain Markdown)
        bibliography: Bibliography text (optional)
        file_path: Path to save the DOCX file
    """
    document = Document()
    
    # --- Title ---
    title_para = document.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # --- Subtitle ---
    if subtitle:
        subtitle_para = document.add_heading(subtitle, level=2)
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # --- Keywords ---
    if keywords:
        kw_text = ", ".join(keywords) if isinstance(keywords, list) else str(keywords)
        p = document.add_paragraph()
        p.add_run("Ключевые слова: ").bold = True
        p.add_run(kw_text)
    
    document.add_paragraph()  # spacing
    
    # --- Main lecture text, paragraph by paragraph ---
    # Split by double newlines (paragraph breaks) but preserve single newlines within paragraphs
    paragraphs = lecture_text.split("\n\n")
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        # Check if it's a heading (starts with #)
        if para_text.startswith('#'):
            # Count heading level
            level = 0
            while level < len(para_text) and para_text[level] == '#':
                level += 1
            
            heading_text = para_text[level:].strip()
            # Limit heading level to 6
            doc_heading = document.add_heading(heading_text, level=min(level, 6))
        else:
            # Regular paragraph - process markdown formatting
            # For multi-line paragraphs, join them with spaces or keep as single paragraph
            # If the text has single newlines, they might be intentional breaks
            # Check if it's multiple lines that should be separate paragraphs or one paragraph
            lines = [line.strip() for line in para_text.split("\n") if line.strip()]
            
            if len(lines) == 1:
                # Single line paragraph
                md_to_docx_paragraph(document, lines[0])
            else:
                # Multiple lines - treat each as a separate paragraph (could be a list or structured text)
                for i, line in enumerate(lines):
                    md_to_docx_paragraph(document, line)
                    # Add spacing between lines if not last
                    if i < len(lines) - 1:
                        document.add_paragraph()  # spacing between lines
    
    # --- Bibliography ---
    if bibliography:
        document.add_page_break()
        document.add_heading("Библиография", level=2)
        
        # Process bibliography line by line
        for line in bibliography.split("\n"):
            line = line.strip()
            if line:
                p = document.add_paragraph(line)
                p.style.font.size = Pt(10)  # Smaller font for bibliography
    
    # Save file
    path = Path(file_path) if file_path else Path("lecture.docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))

